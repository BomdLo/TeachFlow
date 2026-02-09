import streamlit as st
import fitz
import httpx
import json
import re
import asyncio
import sqlite3
import hashlib
from opencc import OpenCC
from docx import Document
import os
import jieba
from wordcloud import WordCloud
import matplotlib.pyplot as plt
from io import BytesIO
import io
import sys
from streamlit_gsheets import GSheetsConnection
import pandas as pd
from datetime import datetime
import secrets
from supabase import create_client, Client
# 建立連線
conn = st.connection("gsheets", type=GSheetsConnection)


# --- 初始化 Supabase ---
@st.cache_resource
def init_supabase():
    
    try:
        url = st.secrets["connections"]["supabase"]["SUPABASE_URL"]
        key = st.secrets["connections"]["supabase"]["SUPABASE_KEY"]
        return create_client(url, key)
    except KeyError as e:
        st.error(f"找不到 Key: {e}。請檢查 Secrets 階層是否正確。")
        st.stop()

# --- 修改後的 Auth 邏輯 ---
def verify_user(username, password):
    try:
        res = supabase.table("users").select("password_hash").eq("username", username).execute()
        if res.data:
            stored_hash = res.data[0]['password_hash']
            return check_hashes(password, stored_hash)
        return False
    except Exception as e:
        st.error(f"DATABASE_ERROR: {e}")
        return False

# --- 修改後的 寫入紀錄 邏輯 ---
def save_task_record(username, task_type, result):
    data = {
        "username": username,
        "task_type": task_type,
        "result": result
    }
    supabase.table("history").insert(data).execute()

# --- 修改後的 讀取紀錄 邏輯 ---
def load_user_history(username):
    res = supabase.table("history").select("*").eq("username", username).order("timestamp", desc=True).limit(5).execute()
    return res.data

def make_hashes(password):
    # 產生一個隨機的「鹽」(Salt)，讓同樣的密碼產生不同的 Hash
    salt = secrets.token_hex(16)
    # 使用 PBKDF2 演算法，並進行 600,000 次疊代，大幅增加暴力破解難度
    key = hashlib.pbkdf2_hmac(
        'sha256', 
        password.encode('utf-8'), 
        salt.encode('utf-8'), 
        600000
    )
    
    # 將鹽與雜湊值存在一起，格式如：salt:hash
    return f"{salt}:{key.hex()}"

def check_hashes(password, hashed_storage):
    try:
        
        # 分離出儲存的鹽與雜湊值
        salt, stored_key = hashed_storage.split(':')
        # 用同樣的鹽對輸入密碼進行計算
        new_key = hashlib.pbkdf2_hmac(
            'sha256', 
            password.encode('utf-8'), 
            salt.encode('utf-8'), 
            600000
        )
        return new_key.hex() == stored_key
    except Exception:
        return False



cc = OpenCC('s2twp')
# --- 1. 設計禁止清單：視覺注入 (CSS Injection) ---
def inject_custom_design():
    st.markdown("""
    <style>
        /* 1. 噪點與非對稱漸變背景 (拒絕純平與紫色) */
        .stApp {
            background-color: #f8fafc;
            background-image: 
                radial-gradient(at 0% 0%, rgba(226, 232, 240, 0.5) 0px, transparent 50%),
                radial-gradient(at 100% 100%, rgba(203, 213, 225, 0.3) 0px, transparent 50%),
                url("https://www.transparenttextures.com/patterns/p6.png");
            background-attachment: fixed;
        }

        /* 2. 專業文字風格 (拒絕 Emoji 功能圖標) */
        h1, h2, h3 {
            font-family: 'Inter', -apple-system, sans-serif;
            color: #1e293b;
            letter-spacing: -0.02em;
        }
        
        /* 3. 按鈕動畫 (拒絕 ease-in-out，改用 Spring 回彈) */
        .stButton>button {
            border-radius: 8px;
            border: 1px solid #e2e8f0;
            background-color: white;
            color: #475569;
            transition: transform 0.2s cubic-bezier(0.34, 1.56, 0.64, 1);
        }
        .stButton>button:hover {
            border-color: #94a3b8;
            color: #1e293b;
            transform: translateY(-1px);
        }
        .stButton>button:active {
            transform: scale(0.98);
        }

        /* 4. 移除容器陰影，改用簡潔邊框 */
        [data-testid="stVerticalBlock"] > div:has(div.stExpander) {
            border: 1px solid #e2e8f0;
            border-radius: 12px;
            background: rgba(255, 255, 255, 0.7);
        }
    </style>
    """, unsafe_allow_html=True)


# --- 2. 初始化 Session State ---
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "username" not in st.session_state:
    st.session_state.username = ""
if "quiz_results" not in st.session_state:
    st.session_state.quiz_results = None
if "display_task" not in st.session_state:
    st.session_state.display_task = None


# --- 3. AI 核心邏輯 (優化 Prompt 以確保 4 個選項) ---
async def run_ai(content, task_type):
    api_url = "https://api.groq.com/openai/v1/chat/completions"
    api_key = st.secrets["GROQ_API_KEY"]
    groq_model = "llama-3.3-70b-versatile"

    if task_type == "生成考題":
        # 這裡加入了 'explanation' 欄位，回擊「NotebookLM 也能做」的質疑
        prompt = (
            "你是一位台灣資深教師。請根據內容出20題『單選題』。\n"
            "要求：1.繁體中文 2.嚴格輸出 JSON 陣列 3.每題 4 個選項。\n"
            "4. 必須包含 'explanation' 欄位，詳述答案理由或課本出處。\n"
            "格式：[{\"question\": \"..\", \"options\": [\"..\",\"..\",\"..\",\"..\"], \"answer\": 0, \"explanation\": \"..\"}]\n\n"
            f"內容：{content[:4000]}"
        )
    else:
        prompt = f"請用繁體中文針對內容進行{task_type}：\n\n{content[:4000]}"

    payload = {
        "model": groq_model,
        "messages": [
            {"role": "system", "content": "TECHNICAL_ASSISTANT_V1: 專業教學助理，語氣精簡。"},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.2
    }

    async with httpx.AsyncClient(timeout=60.0) as client:
        response = await client.post(api_url, headers={"Authorization": f"Bearer {api_key}"}, json=payload)
        return response.json()['choices'][0]['message']['content']

# --- 4. Word 匯出邏輯 ---
def create_docx(quiz_data):
    doc = Document()
    doc.add_heading('TeachFlow 自動生成考卷', 0)
    for i, q in enumerate(quiz_data):
        doc.add_paragraph(f"第 {i + 1} 題：{q['question']}", style='List Number')
        for j, opt in enumerate(q['options']):
            doc.add_paragraph(f"({chr(65 + j)}) {opt}")
    doc.add_page_break()
    doc.add_heading('標準答案', level=1)
    for i, q in enumerate(quiz_data):
        doc.add_paragraph(f"第 {i + 1} 題：({chr(65 + q.get('answer', 0))})")
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


# --- 5. 介面邏輯 ---
st.set_page_config(page_title="TeachFlow AI", layout="wide")
def security_migration_sync(conn_gs):
    """自動偵測並加密 Google Sheets 中的明文密碼"""
    try:
        df = conn_gs.read(ttl=0)
        df.columns = [c.strip() for c in df.columns]
        
        updated = False
        for index, row in df.iterrows():
            pwd = str(row['密碼']).strip()
            
            # 判斷是否為明文：如果沒有 ":" 分隔符號，代表它是 Google Forms 直接寫入的明文
            if ":" not in pwd:
                hashed_pwd = make_hashes(pwd)
                df.at[index, '密碼'] = hashed_pwd
                updated = True
        
        if updated:
            # 寫回 Google Sheets，完成自動加密
            conn_gs.update(data=df)
            st.toast("SECURITY_SYNC: 已自動加密新註冊數據")
            
    except Exception as e:
        # 靜默失敗，不影響登入流程
        pass

def login_ui():
    # 1. 注入全域 CSS
    inject_custom_design()
    
    # 2. 佈局調整
    st.markdown("<br><br>", unsafe_allow_html=True)
    
    # 標題系統感強化
    st.title("TEACHFLOW_AUTH_GATEWAY")
    st.caption("VERSION: 3.0.0_SUPABASE | REGION: TW_EDU")
    
    st.info("SYSTEM_INFO: 已遷移至 PostgreSQL 隔離架構，支援高併發存取。")
    
    # 初始化 Supabase (假設你已定義 init_supabase 函數)
    supabase = init_supabase()
    
    # 使用 Tabs
    tab1, tab2 = st.tabs(["SIGN_IN", "REGISTRATION"])

    with tab2:
        st.markdown("### ACCOUNT_REGISTRATION")
        st.write("目前註冊由 Supabase 安全驗證層接管。")
        with st.container(border=True):
            reg_user = st.text_input("SET_ID_ACCOUNT", placeholder="欲註冊的帳號")
            reg_pass = st.text_input("SET_ACCESS_PASSWORD", type='password', placeholder="欲設定的密碼")
            if st.button("EXECUTE_REGISTRATION", use_container_width=True):
                if reg_user and reg_pass:
                    # 密碼雜湊處理
                    hashed_pw = make_hashes(reg_pass)
                    try:
                        # 寫入 Supabase user 表單
                        supabase.table("users").insert({
                            "username": reg_user.strip(),
                            "password_hash": hashed_pw
                        }).execute()
                        st.success("SUCCESS: 帳號已建立，請切換至登入頁面。")
                    except Exception as e:
                        st.error(f"REG_ERROR: 帳號可能已存在或系統異常")
                else:
                    st.warning("FIELD_REQUIRED: 請填寫完整資訊")

    with tab1:
        with st.container(border=True):
            user_input = st.text_input("ID_ACCOUNT", placeholder="輸入註冊帳號")
            pass_input = st.text_input("ACCESS_PASSWORD", type='password', placeholder="輸入安全密碼")
            
            if st.button("VERIFY_AND_LOGIN", use_container_width=True):
                if user_input and pass_input:
                    try:
                        # 步驟 A: 直接從 Supabase 撈取該使用者的雜湊密碼 (不再讀取全表)
                        search_id = str(user_input).strip()
                        response = supabase.table("users").select("password_hash").eq("username", search_id).execute()
                        
                        if response.data:
                            stored_hash = response.data[0]['password_hash']
                            
                            # 步驟 B: 安全比對
                            if check_hashes(pass_input.strip(), stored_hash):
                                st.session_state.logged_in = True
                                st.session_state.username = search_id
                                st.success("AUTH_SUCCESS: 正在跳轉工作站...")
                                st.rerun()
                            else:
                                st.error("AUTH_ERROR: 密碼驗證失敗")
                        else:
                            st.error("AUTH_ERROR: 找不到該帳號")
                            
                    except Exception as e:
                        st.error(f"SYSTEM_ERROR: 連線資料庫失敗")
                else:
                    st.warning("FIELD_REQUIRED: 帳號密碼不可為空")
            # --- 關鍵字雲生成邏輯 ---
def generate_wordcloud(text):
    # 1. 斷詞處理
    # 建議加入一些自定義的停止詞 (Stopwords)，過濾掉「的」、「是」、「在」等無意義字
    words = jieba.cut(text)
    clean_text = " ".join([word for word in words if len(word) > 1])

    # 2. 設定字型路徑 (請根據你的作業系統修改路徑)
    # Mac 範例路徑
    font_path = "NotoSansTC-VariableFont_wght.ttf"  # 確保檔案跟 app.py 放一起
    # Windows 範例路徑: "C:/Windows/Fonts/msjh.ttc"

    # 3. 建立文字雲物件
    wc = WordCloud(
        font_path=font_path,
        background_color="white",
        width=800,
        height=400,
        max_words=100,
        colormap="viridis"  # 顏色主題
    )

    # 4. 產生圖片
    wc.generate(clean_text)

    # 5. 將圖片轉為 Streamlit 可讀取的格式
    img_buffer = BytesIO()
    plt.figure(figsize=(10, 5))
    plt.imshow(wc, interpolation='bilinear')
    plt.axis("off")
    plt.tight_layout(pad=0)
    plt.savefig(img_buffer, format='png')
    return img_buffer

def main_app():
    # 注入視覺設計
    inject_custom_design()
    
    # 初始化 Supabase 用戶端 (確保 init_supabase 已定義)
    supabase = init_supabase()

    st.title("TEACHFLOW_WORKSPACE_V3")
    st.caption(f"ACTIVE_USER: {st.session_state.username} | DB_ENGINE: POSTGRESQL (SUPABASE)")

    # --- 1. 側邊欄：從 Supabase 讀取紀錄 ---
    with st.sidebar:
        st.markdown("### SYSTEM_CONTROL")
        model_name = st.selectbox("MODEL_SELECT", ["llama-3.3-70b", "deepseek-r1:7b"])
        if st.button("LOGOUT_SESSION"):
            st.session_state.logged_in = False
            st.rerun()

        st.divider()
        st.markdown("### DATA_HISTORY")
        
        try:
            # 💡 核心優化：直接從 Supabase 查詢該使用者的前 5 筆紀錄
            response = supabase.table("history") \
                .select("*") \
                .eq("username", str(st.session_state.username)) \
                .order("timestamp", desc=True) \
                .limit(5) \
                .execute()
            
            records = response.data
            if records:
                for i, row in enumerate(records):
                    # 格式化時間標籤
                    time_label = row['timestamp'][5:16].replace('T', ' ')
                    if st.button(f"REC_{time_label} | {row['task_type']}", key=f"hist_{i}", use_container_width=True):
                        st.session_state.quiz_results = row['result']
                        st.session_state.display_task = row['task_type']
                        st.rerun()
            else:
                st.caption("NO_RECORDS_AVAILABLE")
        except Exception as e:
            st.caption("DATABASE_CONNECTION_PENDING")

    # --- 2. 主畫面佈局 ---
    col_meta, col_workspace = st.columns([1, 2.5], gap="large")

    with col_meta:
        st.markdown("### 01_SOURCE_UPLOAD")
        uploaded_file = st.file_uploader("UPLOAD_PDF", type="pdf", label_visibility="collapsed")
        
        if uploaded_file:
            doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
            full_text = "".join([page.get_text() for page in doc])
            st.code(f"METRICS: {len(full_text)} CHARS", language="bash")
            
            st.markdown("### 02_TASK_CONFIGURATION")
            task = st.radio("SELECT_OPERATION", ["重點摘要", "生成考題", "教學策略建議"], label_visibility="collapsed")
            
            if st.button("EXECUTE_AI_ANALYSIS", use_container_width=True):
                with st.spinner("AI_THINKING..."):
                    raw = asyncio.run(run_ai(full_text, task))
                    processed = cc.convert(raw).replace("後-end", "後端")
                    processed = re.sub(r'<think>.*?</think>', '', processed, flags=re.DOTALL)
                    processed = re.sub(r'```json|```', '', processed)
                    
                    # --- 💡 寫入 Supabase 邏輯 (取代原本複雜的 pd.concat) ---
                    try:
                        new_data = {
                            "username": st.session_state.username,
                            "task_type": task,
                            "result": processed
                            # timestamp 由 Supabase 自動生成 (DEFAULT now())
                        }
                        supabase.table("history").insert(new_data).execute()
                        st.toast("✅ 紀錄已同步至 Supabase")
                    except Exception as e:
                        st.error(f"SUPABASE_SYNC_ERROR: {str(e)}")
                    
                    st.session_state.quiz_results = processed
                    st.session_state.display_task = task
                    st.rerun()

            if st.button("GENERATE_WORD_CLOUD", use_container_width=True):
                with st.spinner("ANALYZING..."):
                    cloud_img = generate_wordcloud(full_text)
                    st.session_state.current_cloud = cloud_img

    with col_workspace:
        if "current_cloud" in st.session_state:
            st.image(st.session_state.current_cloud, use_container_width=True)

        if st.session_state.quiz_results:
            st.markdown(f"### 03_OUTPUT: {st.session_state.display_task}")
            res = st.session_state.quiz_results
            
            if st.session_state.display_task == "生成考題":
                json_match = re.search(r'\[.*\]', res, re.DOTALL)
                if json_match:
                    try:
                        quiz_data = json.loads(json_match.group())
                        for i, q in enumerate(quiz_data):
                            with st.container(border=True):
                                st.markdown(f"**Q{i + 1}: {q['question']}**")
                                st.radio("OPTIONS", q['options'], key=f"q_{i}_{hash(res)}", label_visibility="collapsed")
                                with st.expander("VIEW_LOGIC"):
                                    st.markdown(f"**CORRECT:** {q['options'][q.get('answer', 0)]}")
                                    if 'explanation' in q:
                                        st.caption(f"LOGIC: {q['explanation']}")
                        
                        st.download_button("DOWNLOAD_DOCX", create_docx(quiz_data), "exam.docx", use_container_width=True)
                    except Exception as e:
                        st.error("JSON_PARSE_ERROR: AI 回傳格式異常")
                        st.text_area("RAW_DATA", res, height=200)
                else:
                    st.text_area("RAW_OUTPUT", res, height=400)
            else:
                st.markdown(res)
        else:
            st.info("AWAITING_INPUT: 請上傳檔案並執行 AI 分析。")

    st.divider()
    st.caption("USER_FEEDBACK_REQUIRED")
    st.link_button("SUBMIT_FEEDBACK", "https://forms.gle/p9iJdyMYaZBg9NxMA")



if not st.session_state.logged_in:
    login_ui()
else:
    main_app()
